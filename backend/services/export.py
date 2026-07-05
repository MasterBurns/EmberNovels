import re
import tempfile
import markdown
import os
from pathlib import Path
from typing import List, Dict, Any, Tuple
from datetime import datetime

# Optional imports handled safely
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
except ImportError:
    SimpleDocTemplate = None

try:
    import ebooklib
    from ebooklib import epub
except ImportError:
    epub = None

from backend.services.storage import StorageService

class ExportService:
    @classmethod
    def get_project_chapters_content(cls, project_id: str, chapter_ids: List[str] = None) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
        """Load project metadata and selected or all active chapters with content."""
        project = StorageService.get_project_metadata(project_id)
        if not project:
            raise ValueError("Projekt nicht gefunden.")
            
        chapters_meta = StorageService.list_chapters(project_id)
        chapters = []
        
        for ch_meta in chapters_meta:
            # Filter by chapter_ids if specified
            if chapter_ids is not None and ch_meta['id'] not in chapter_ids:
                continue
                
            ch_data = StorageService.get_chapter_content(project_id, ch_meta['id'])
            if ch_data and "error" not in ch_data:
                chapters.append({
                    "id": ch_meta['id'],
                    "title": ch_meta['title'],
                    "content": ch_data.get('content', '')
                })
        return project, chapters

    @classmethod
    def export_project(cls, project_id: str, file_format: str, chapter_ids: List[str] = None) -> Tuple[Path, str]:
        """
        Export project chapters into a single file of the given format.
        Returns: Tuple[Path_to_temp_file, File_name]
        """
        project, chapters = cls.get_project_chapters_content(project_id, chapter_ids)
        if not chapters:
            raise ValueError("Keine Kapitel für den Export vorhanden oder ausgewählt.")
            
        project_title = project.get('title', 'Unbenanntes Buch')
        project_desc = project.get('description', '')
        
        # Sanitize filename for download header
        safe_title = re.sub(r'[^a-zA-Z0-9_-]', '_', project_title)
        temp_dir = Path(tempfile.gettempdir())
        
        if file_format == 'txt':
            filename = f"{safe_title}.txt"
            file_path = temp_dir / filename
            cls._export_txt(file_path, project_title, project_desc, chapters)
            
        elif file_format == 'html':
            filename = f"{safe_title}.html"
            file_path = temp_dir / filename
            cls._export_html(file_path, project_title, project_desc, chapters)
            
        elif file_format == 'docx':
            if not Document:
                raise ImportError("python-docx ist nicht installiert.")
            filename = f"{safe_title}.docx"
            file_path = temp_dir / filename
            cls._export_docx(file_path, project_title, project_desc, chapters)
            
        if file_format == 'pdf':
            if not SimpleDocTemplate:
                raise ImportError("reportlab ist nicht installiert.")
            filename = f"{safe_title}.pdf"
            file_path = temp_dir / filename
            cls._export_pdf(file_path, project_id, project, chapters)
            
        elif file_format == 'epub':
            if not epub:
                raise ImportError("ebooklib ist nicht installiert.")
            filename = f"{safe_title}.epub"
            file_path = temp_dir / filename
            cls._export_epub(file_path, project_id, project, chapters)
            
        elif file_format == 'zip':
            filename = f"{safe_title}_Backup.zip"
            file_path = temp_dir / filename
            cls._export_zip(file_path, project_id)
            
        else:
            raise ValueError(f"Ungültiges Format: {file_format}")
            
        return file_path, filename

    @classmethod
    def _export_zip(cls, file_path: Path, project_id: str):
        import zipfile
        project_dir = StorageService.get_projects_dir() / project_id
        if not project_dir.exists():
            raise ValueError("Projektverzeichnis nicht gefunden.")
            
        with zipfile.ZipFile(file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(project_dir):
                for file in files:
                    file_p = Path(root) / file
                    # Calculate arcname relative to the projects_dir so it unzips neatly as `project_id/...`
                    arcname = file_p.relative_to(StorageService.get_projects_dir())
                    zipf.write(file_p, arcname)

    @classmethod
    def _export_txt(cls, file_path: Path, title: str, desc: str, chapters: List[Dict[str, Any]]):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"=========================================\n")
            f.write(f"{title}\n")
            if desc:
                f.write(f"{desc}\n")
            f.write(f"=========================================\n\n")
            
            for ch in chapters:
                f.write(f"--- {ch['title']} ---\n\n")
                f.write(ch['content'])
                f.write("\n\n")

    @classmethod
    def _export_html(cls, file_path: Path, title: str, desc: str, chapters: List[Dict[str, Any]]):
        body_content = ""
        toc_content = "<ul>"
        
        for i, ch in enumerate(chapters):
            ch_html = markdown.markdown(ch['content'])
            anchor = f"chapter-{i}"
            toc_content += f'<li><a href="#{anchor}">{ch["title"]}</a></li>'
            body_content += f'<section id="{anchor}" style="margin-top: 48px; border-top: 1px solid #ddd; padding-top: 24px;">'
            body_content += f'<h1>{ch["title"]}</h1>'
            body_content += ch_html
            body_content += '</section>'
            
        toc_content += "</ul>"
        
        html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Georgia', serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 40px auto;
            padding: 0 20px;
        }}
        h1, h2, h3 {{ font-family: sans-serif; }}
        nav {{ background: #f9f9f9; padding: 16px; border-radius: 8px; margin-bottom: 40px; }}
        .cover {{ text-align: center; margin-bottom: 60px; }}
        .cover h1 {{ font-size: 3em; margin-bottom: 10px; }}
        .cover p {{ font-style: italic; color: #666; font-size: 1.2em; }}
    </style>
</head>
<body>
    <div class="cover">
        <h1>{title}</h1>
        <p>{desc}</p>
    </div>
    <nav>
        <h2>Inhaltsverzeichnis</h2>
        {toc_content}
    </nav>
    {body_content}
</body>
</html>"""
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_template)

    @classmethod
    def _export_docx(cls, file_path: Path, title: str, desc: str, chapters: List[Dict[str, Any]]):
        doc = Document()
        doc.add_heading(title, 0)
        if desc:
            p = doc.add_paragraph()
            p.add_run(desc).italic = True
            
        doc.add_page_break()
        
        for ch in chapters:
            doc.add_heading(ch['title'], level=1)
            lines = ch['content'].split('\n')
            for line in lines:
                if line.strip().startswith('#'):
                    level = min(5, len(line) - len(line.lstrip('#')))
                    doc.add_heading(line.strip().lstrip('#').strip(), level=level + 1)
                else:
                    doc.add_paragraph(line)
            doc.add_page_break()
            
        doc.save(str(file_path))

    @classmethod
    def _html_to_flowables(cls, html_content: str, styles) -> List[Any]:
        from html.parser import HTMLParser
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.styles import ParagraphStyle
        
        class PDFHTMLParser(HTMLParser):
            def __init__(self):
                super().__init__()
                self.flowables = []
                self.current_tag = None
                self.current_text = ""
                self.list_stack = []
                
            def handle_starttag(self, tag, attrs):
                if tag in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote']:
                    self.current_tag = tag
                    self.current_text = ""
                elif tag in ['ul', 'ol']:
                    self.list_stack.append(tag)
                elif tag in ['strong', 'em', 'b', 'i', 'code', 'pre', 'a', 'span']:
                    mapped_tag = tag
                    if tag == 'strong': mapped_tag = 'b'
                    elif tag == 'em': mapped_tag = 'i'
                    attr_str = "".join([f' {k}="{v}"' for k, v in attrs])
                    self.current_text += f"<{mapped_tag}{attr_str}>"
                    
            def handle_endtag(self, tag):
                if tag in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote']:
                    style_name = 'BookBody'
                    if tag == 'h1': style_name = 'Heading1'
                    elif tag == 'h2': style_name = 'Heading2'
                    elif tag == 'h3': style_name = 'Heading3'
                    elif tag == 'h4' or tag == 'h5' or tag == 'h6': style_name = 'Heading4'
                    elif tag == 'li': style_name = 'Bullet'
                    
                    style = styles.get(style_name, styles['Normal'])
                    if tag == 'blockquote':
                        bq_style = ParagraphStyle(
                            'BlockQuote',
                            parent=styles['Normal'],
                            leftIndent=20,
                            fontName='Helvetica-Oblique',
                            spaceAfter=10
                        )
                        style = bq_style
                        
                    p_text = self.current_text.strip()
                    if p_text:
                        self.flowables.append(Paragraph(p_text, style))
                        if tag.startswith('h'):
                            self.flowables.append(Spacer(1, 8))
                        else:
                            self.flowables.append(Spacer(1, 10))
                    self.current_tag = None
                    self.current_text = ""
                elif tag in ['ul', 'ol']:
                    if self.list_stack:
                        self.list_stack.pop()
                elif tag in ['strong', 'em', 'b', 'i', 'code', 'pre', 'a', 'span']:
                    mapped_tag = tag
                    if tag == 'strong': mapped_tag = 'b'
                    elif tag == 'em': mapped_tag = 'i'
                    self.current_text += f"</{mapped_tag}>"
                    
            def handle_data(self, data):
                if self.current_tag is not None:
                    escaped_data = data.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    self.current_text += escaped_data
                    
        parser = PDFHTMLParser()
        parser.feed(html_content)
        return parser.flowables

    @classmethod
    def _export_pdf(cls, file_path: Path, project_id: str, project: Dict[str, Any], chapters: List[Dict[str, Any]]):
        from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, NextPageTemplate
        from reportlab.lib.pagesizes import landscape
        
        styles = getSampleStyleSheet()
        
        title = project.get('title', 'Unbenanntes Buch')
        author = project.get('author', 'EmberNovels Author')
        desc = project.get('description', '')
        imprint = project.get('imprint', f"© {datetime.now().year} {author}. Alle Rechte vorbehalten.")
        
        title_style = ParagraphStyle(
            'CoverTitle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=28, leading=34, alignment=TA_CENTER, spaceAfter=15
        )
        author_style = ParagraphStyle(
            'CoverAuthor', parent=styles['Normal'], fontName='Helvetica', fontSize=16, leading=20, alignment=TA_CENTER, spaceBefore=20, spaceAfter=30
        )
        desc_style = ParagraphStyle(
            'CoverDesc', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=14, leading=18, alignment=TA_CENTER, spaceAfter=30
        )
        imprint_style = ParagraphStyle(
            'Imprint', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14, alignment=TA_CENTER, textColor='#555555'
        )
        h1_style = ParagraphStyle(
            'ChapterHeader', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=20, leading=24, spaceBefore=20, spaceAfter=15, keepWithNext=True, alignment=TA_CENTER
        )
        body_style = ParagraphStyle(
            'BookBody', parent=styles['BodyText'], fontName='Helvetica', fontSize=11, leading=16, alignment=TA_JUSTIFY, firstLineIndent=20, spaceAfter=5
        )
        
        doc = BaseDocTemplate(str(file_path), pagesize=letter)
        
        frame_portrait = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='portrait_frame')
        # Landscape frame flips width and height
        frame_landscape = Frame(doc.leftMargin, doc.bottomMargin, doc.height, doc.width, id='landscape_frame')
        
        template_portrait = PageTemplate(id='portrait', frames=[frame_portrait], pagesize=letter)
        template_landscape = PageTemplate(id='landscape', frames=[frame_landscape], pagesize=landscape(letter))
        
        doc.addPageTemplates([template_portrait, template_landscape])
        
        story = []
        
        # Cover Image
        if project.get('has_cover'):
            cover_path = StorageService.get_projects_dir() / project_id / "cover.jpg"
            if cover_path.exists():
                try:
                    img = Image(str(cover_path))
                    img.drawHeight = 500
                    img.drawWidth = 500 * (img.imageWidth / img.imageHeight)
                    story.append(img)
                    story.append(PageBreak())
                except Exception as e:
                    pass
        
        # Title Page
        story.append(Spacer(1, 150))
        story.append(Paragraph(title, title_style))
        if author:
            story.append(Paragraph(f"von {author}", author_style))
        if desc:
            story.append(Paragraph(desc, desc_style))
        story.append(PageBreak())
        
        # Imprint Page
        if imprint:
            story.append(Spacer(1, 300))
            imprint_html = imprint.replace('\\n', '<br/>').replace('\n', '<br/>')
            story.append(Paragraph(imprint_html, imprint_style))
            story.append(PageBreak())
            
        # Get chapter metadata mapping
        chapters_metadata = project.get('chapters_metadata', {})
        
        # Chapters
        chapter_counter = 0
        for ch in chapters:
            ch_meta = chapters_metadata.get(ch['id'], {})
            ch_type = ch_meta.get('type', 'standard')
            
            if ch_type == 'map':
                story.append(NextPageTemplate('landscape'))
                story.append(PageBreak())
            else:
                story.append(NextPageTemplate('portrait'))
            
            if not ch_type or ch_type == 'standard':
                chapter_counter += 1
                ch_title = f"{chapter_counter}. {ch['title']}"
            else:
                ch_title = ch['title']
                
            story.append(Paragraph(ch_title, h1_style))
            
            # Map chapter handling (make images fit landscape)
            ch_html = markdown.markdown(ch['content'])
            if ch_type == 'map':
                # Parse HTML but try to scale images to fit landscape
                ch_flowables = cls._html_to_flowables(ch_html, styles)
                for f in ch_flowables:
                    if isinstance(f, Image):
                        # Landscape width is doc.height
                        max_w = doc.height
                        max_h = doc.width - 100
                        aspect = f.imageWidth / f.imageHeight
                        f.drawWidth = min(f.imageWidth, max_w)
                        f.drawHeight = f.drawWidth / aspect
                        if f.drawHeight > max_h:
                            f.drawHeight = max_h
                            f.drawWidth = f.drawHeight * aspect
                story.extend(ch_flowables)
            else:
                story.extend(cls._html_to_flowables(ch_html, styles))
                
            story.append(PageBreak())
            story.append(NextPageTemplate('portrait'))
            
        doc.build(story)

    @classmethod
    def _export_epub(cls, file_path: Path, project_id: str, project: Dict[str, Any], chapters: List[Dict[str, Any]]):
        book = epub.EpubBook()
        book.set_identifier(project_id)
        
        title = project.get('title', 'Unbenanntes Buch')
        author = project.get('author', 'EmberNovels Author')
        desc = project.get('description', '')
        imprint = project.get('imprint', f"© {datetime.now().year} {author}. Alle Rechte vorbehalten.")
        
        book.set_title(title)
        book.set_language('de')
        book.add_author(author)
        
        # Add basic CSS
        style = '''
            body { font-family: "Georgia", serif; text-align: justify; line-height: 1.6; }
            h1 { text-align: center; margin-top: 2em; margin-bottom: 2em; page-break-before: always; }
            h2, h3 { text-align: center; margin-top: 1.5em; margin-bottom: 1em; }
            p { margin: 0; text-indent: 1.5em; }
            p.no-indent { text-indent: 0; }
            .center { text-align: center; text-indent: 0; }
            .imprint { font-size: 0.9em; text-align: center; margin-top: 50%; color: #555; text-indent: 0; }
        '''
        nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)
        
        spine = []
        
        # Handle Cover Image
        if project.get('has_cover'):
            cover_path = StorageService.get_projects_dir() / project_id / "cover.jpg"
            if cover_path.exists():
                with open(cover_path, 'rb') as f:
                    book.set_cover("cover.jpg", f.read())
                # Spine will automatically get the cover page
                spine.append('cover')
        
        # Title Page
        title_page = epub.EpubHtml(title='Titelblatt', file_name='title.xhtml', lang='de')
        title_page.content = f'<div class="center" style="margin-top: 100px;"><h1>{title}</h1>'
        if author:
            title_page.content += f'<h3 style="margin-top: 50px;">von {author}</h3>'
        if desc:
            title_page.content += f'<p class="center" style="margin-top: 30px;"><i>{desc}</i></p>'
        title_page.content += '</div>'
        title_page.add_item(nav_css)
        book.add_item(title_page)
        spine.append(title_page)
        
        # Imprint Page
        if imprint:
            imprint_page = epub.EpubHtml(title='Impressum', file_name='imprint.xhtml', lang='de')
            imprint_html = imprint.replace('\\n', '<br>').replace('\n', '<br>')
            imprint_page.content = f'<div class="imprint">{imprint_html}</div>'
            imprint_page.add_item(nav_css)
            book.add_item(imprint_page)
            spine.append(imprint_page)
            
        chapters_metadata = project.get('chapters_metadata', {})
        
        # Chapters
        toc = []
        chapter_counter = 0
        for i, ch in enumerate(chapters):
            ch_meta = chapters_metadata.get(ch['id'], {})
            ch_type = ch_meta.get('type', 'standard')
            
            if not ch_type or ch_type == 'standard':
                chapter_counter += 1
                ch_title = f"{chapter_counter}. {ch['title']}"
            else:
                ch_title = ch['title']
                
            ch_html = markdown.markdown(ch['content'])
            # Basic attempt to add no-indent class to first paragraph
            ch_html = ch_html.replace('<p>', '<p class="no-indent">', 1)
            
            c = epub.EpubHtml(title=ch_title, file_name=f'chap_{i}.xhtml', lang='de')
            c.content = f"<h1>{ch_title}</h1>\n{ch_html}"
            c.add_item(nav_css)
            book.add_item(c)
            spine.append(c)
            toc.append(c)
            
        book.toc = tuple(toc) # Only put chapters in TOC
        book.spine = ['nav'] + spine
        book.add_item(epub.EpubNav())
        book.add_item(epub.EpubNcx())
        
        epub.write_epub(str(file_path), book, {})
